import unittest
import tempfile
import os
from docx import Document as DocxDocument
from free_vigilance_reduction.documents.docx_processor import DocxProcessor


class TestDocxProcessor(unittest.TestCase):
    def setUp(self):
        self.temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc = DocxDocument()
        doc.add_paragraph("Это тестовый текст.")
        doc.save(self.temp_file.name)
        self.processor = DocxProcessor(self.temp_file.name)

    def tearDown(self):
        os.unlink(self.temp_file.name)
        if os.path.exists(self.processor.metadata.get('redacted_path', '')):
            os.unlink(self.processor.metadata['redacted_path'])

    def test_get_text(self):
        text = self.processor.get_text()
        self.assertEqual(text.strip(), "Это тестовый текст.")

    def test_create_redacted_copy(self):
        redacted_text = "[REDACTED TEXT]"
        self.processor.create_redacted_copy(redacted_text)
        redacted_path = self.processor.metadata['redacted_path']
        self.assertTrue(os.path.exists(redacted_path))

        redacted_doc = DocxDocument(redacted_path)
        content = "\n".join(p.text for p in redacted_doc.paragraphs)
        self.assertEqual(content.strip(), redacted_text)


if __name__ == '__main__':
    unittest.main()
