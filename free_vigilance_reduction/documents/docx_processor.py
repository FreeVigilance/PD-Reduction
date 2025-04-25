"""
Обработчик DOCX-документов.
"""

from .base import Document
from pathlib import Path
from docx import Document as DocxDocument


class DocxProcessor(Document):
    """
    Класс для извлечения текста из DOCX-документов и создания редактированной копии.
    """

    def get_text(self) -> str:
        """
        Извлекает текст из .docx файла.

        Returns:
            str: Извлечённый текст.
        """
        docx_now = DocxDocument(self.file_path)
        paragraphs_now = []

        for paragraph_now in docx_now.paragraphs:
            if paragraph_now.text.strip():
                paragraphs_now.append(paragraph_now.text)

        self.text_content = "\n".join(paragraphs_now)
        return self.text_content

    def create_redacted_copy(self, reduced_text: str) -> None:
        """
        Создаёт редактированную копию .docx документа.

        Args:
            reduced_text (str): Текст после анонимизации.
        """
        new_docx_now = DocxDocument()
        lines_now = reduced_text.split("\n")

        for line_now in lines_now:
            new_docx_now.add_paragraph(line_now)

        redacted_path_now = Path(self.file_path).with_name(
            Path(self.file_path).stem + "_redacted.docx"
        )
        new_docx_now.save(redacted_path_now)
        self.metadata['redacted_path'] = str(redacted_path_now)
